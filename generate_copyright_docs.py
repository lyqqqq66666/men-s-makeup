import os
import re
import math
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 配置信息
SOFTWARE_NAME = "智颜方正—智能图像矫正与风格男妆生成系统"
VERSION = "V1.0"
SRC_ROOT = "/Users/LYQ/Desktop/大三资料/数字图像处理/智颜方正—智能图像矫正与风格男妆生成系统/src"
OUTPUT_DIR = "/Users/LYQ/Desktop/大三资料/数字图像处理/智颜方正—智能图像矫正与风格男妆生成系统"

def clean_code(code):
    # 移除 /* */ 注释
    code = re.sub(r'/\*.*?\*/', '', code, flags=re.DOTALL)
    # 移除 // 注释
    code = re.sub(r'//.*', '', code)
    # 移除空行
    lines = [line.strip() for line in code.splitlines() if line.strip()]
    return lines

def get_all_src_files(root):
    src_files = []
    # 按照优先级排序：vue > ts > js > css/html
    ext_priority = {'.vue': 0, '.ts': 1, '.js': 2, '.html': 3, '.css': 4}
    
    for dirpath, _, filenames in os.walk(root):
        if 'node_modules' in dirpath:
            continue
        for filename in filenames:
            ext = os.path.splitext(filename)[1]
            if ext in ext_priority:
                src_files.append((ext_priority[ext], os.path.join(dirpath, filename)))
    
    src_files.sort()
    return [f[1] for f in src_files]

def generate_code_doc(all_lines, output_path):
    doc = Document()
    # 设置页眉
    header = doc.sections[0].header
    p = header.paragraphs[0]
    p.text = f"{SOFTWARE_NAME} {VERSION}"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 合规性处理：前30页 + 后30页
    # 每页不少于50行，60页大约需要 3000 行
    target_lines = all_lines
    if len(all_lines) > 3000:
        front_30 = all_lines[:1500]
        back_30 = all_lines[-1500:]
        target_lines = front_30 + back_30
    
    # 写入文档
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    
    code_text = "\n".join(target_lines)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9) # 小号字体确保每页行数
    
    doc.save(output_path)
    print(f"已生成源代码鉴别材料: {output_path}")

def generate_manual(output_path):
    doc = Document()
    doc.add_heading('用户使用说明书', 0)
    
    doc.add_heading('1. 引言', 1)
    doc.add_paragraph('智颜方正—智能图像矫正与风格男妆生成系统是一款基于计算机视觉技术的智能图像处理工具。')
    
    doc.add_heading('2. 软件概述', 1)
    doc.add_paragraph('系统集成了图像矫正、人脸检测、风格妆容迁移等核心功能。')
    
    doc.add_heading('3. 软件功能操作', 1)
    
    functions = [
        ("用户登录/注册", "进入系统首页进行身份验证。"),
        ("图像上传与预览", "用户可以点击上传区域，选择需要处理的原始照片。"),
        ("智能图像矫正", "系统自动检测并修正照片的倾斜、光照等缺陷。"),
        ("风格男妆选择", "提供多种不同风格的男性妆容效果（如自然、商务、硬朗等）。"),
        ("生成与下载", "一键生成预览结果，并支持高清图像下载保存。")
    ]
    
    for title, desc in functions:
        doc.add_heading(f'3.{functions.index((title, desc))+1} {title}', 2)
        doc.add_paragraph(desc)
        doc.add_paragraph(f'[请在此插入截图：{title}界面]')
        doc.add_paragraph("\n")
    
    doc.save(output_path)
    print(f"已生成用户使用说明书大纲: {output_path}")

def generate_info(total_lines, output_path):
    doc = Document()
    doc.add_heading('申请表关键信息', 0)
    
    info = {
        "软件全称": SOFTWARE_NAME,
        "软件简称": "智颜方正",
        "版本号": VERSION,
        "总行数": len(total_lines),
        "主要功能特点": [
            "基于深度学习的智能图像自动矫正算法。",
            "高精度的男性面部特征点检测与妆容实时迁移。",
            "支持多风格、个性化的男妆效果生成。",
            "响应式 Web 界面，支持多终端平滑访问。"
        ],
        "技术特点": [
            "采用 Vue 3 + TypeScript 响应式架构。",
            "集成深度神经网络模型进行端到端图像转换。",
            "使用 Vite 构建系统，具备极速的热重载和构建性能。",
            "前后端分离，通过 RESTful API 进行高效数据交换。"
        ],
        "硬件环境": "CPU 2.0GHz, 16GB RAM, 支持主流显卡加速",
        "软件环境": "Windows/macOS, Node.js 18+, Chrome 100+"
    }
    
    for key, value in info.items():
        if isinstance(value, list):
            doc.add_heading(key, 1)
            for item in value:
                doc.add_paragraph(item, style='List Bullet')
        else:
            doc.add_paragraph(f"{key}: {value}")
            
    doc.save(output_path)
    print(f"已生成申请表关键信息: {output_path}")

if __name__ == "__main__":
    src_files = get_all_src_files(SRC_ROOT)
    all_cleaned_lines = []
    
    for f_path in src_files:
        try:
            with open(f_path, 'r', encoding='utf-8') as f:
                content = f.read()
                all_cleaned_lines.extend(clean_code(content))
        except Exception as e:
            print(f"读取文件失败 {f_path}: {e}")
            
    generate_code_doc(all_cleaned_lines, os.path.join(OUTPUT_DIR, "源代码鉴别材料.docx"))
    generate_manual(os.path.join(OUTPUT_DIR, "用户使用说明书.docx"))
    generate_info(all_cleaned_lines, os.path.join(OUTPUT_DIR, "申请表关键信息.docx"))
